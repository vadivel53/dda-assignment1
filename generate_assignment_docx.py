from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_heading(document, text, level=1):
    document.add_heading(text, level=level)


def add_para(document, text, bold=False):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = bold
    return p


def add_code_block(document, code_text):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def set_document_styles(document):
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)


def add_page_number_footer(document):
    footer = document.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = p.add_run()
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def format_all_paragraphs(document):
    for p in document.paragraphs:
        if not p.text.strip():
            continue
        if p.style.name.startswith("Heading"):
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            continue
        if any(r.font.name == "Consolas" for r in p.runs):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            continue
        if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)


doc = Document()
set_document_styles(doc)
add_page_number_footer(doc)

# Cover Page
title = doc.add_heading("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

campus = doc.add_paragraph("Work Integrated Learning Programmes Division")
campus.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

course = doc.add_paragraph("Course: Database Design and Applications (CSIZG518/SEZG518/SSZG518)")
course.alignment = WD_ALIGN_PARAGRAPH.CENTER

assign_title = doc.add_paragraph("Assignment: Assignment-1 - Problem Statement Set 03")
assign_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")
doc.add_paragraph("")

student = doc.add_paragraph()
student.alignment = WD_ALIGN_PARAGRAPH.CENTER
student.add_run("Submitted By\n").bold = True
student.add_run("Name: VADIVEL PERIYASAMY\n")
student.add_run("Roll No: 2025TM93235\n")

submit_date = doc.add_paragraph(f"Date: {date.today().strftime('%d-%m-%Y')}")
submit_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

main_title = doc.add_heading("Complete Assignment Solution", level=1)
main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

add_heading(doc, "Q1. Conceptual Schema and ER Design for Notown Records", level=1)

add_heading(doc, "Entities and Attributes", level=2)
add_para(doc, "1. MUSICIAN(ssn [PK], name, address, phone)")
add_para(doc, "2. INSTRUMENT(inst_id [PK], inst_name, musical_key)")
add_para(doc, "3. ALBUM(album_id [PK], title, copyright_date, format, album_identifier [UNIQUE])")
add_para(doc, "4. SONG(song_id [PK], title, author)")

add_heading(doc, "Relationships and Cardinalities", level=2)
add_para(doc, "1. PLAYS: MUSICIAN M:N INSTRUMENT")
add_para(doc, "2. PERFORMS: MUSICIAN M:N SONG (total participation of SONG)")
add_para(doc, "3. CONTAINS: ALBUM 1:N SONG (each SONG appears in exactly one ALBUM)")
add_para(doc, "4. PRODUCES: MUSICIAN 1:N ALBUM (each ALBUM has exactly one producer)")

add_heading(doc, "Clear ER Diagram (Text Notation)", level=2)
add_code_block(
    doc,
    "Entities:\n"
    "  MUSICIAN(ssn PK, name, address, phone)\n"
    "  INSTRUMENT(inst_id PK, inst_name, musical_key)\n"
    "  SONG(song_id PK, title, author)\n"
    "  ALBUM(album_id PK, title, copyright_date, format, album_identifier UNIQUE)\n"
    "\n"
    "Relationships (Crow's-foot style):\n"
    "  MUSICIAN  M:N  INSTRUMENT   [PLAYS]\n"
    "  MUSICIAN  M:N  SONG         [PERFORMS]    (SONG has total participation: at least 1 musician)\n"
    "  ALBUM     1:N  SONG         [CONTAINS]    (each SONG belongs to exactly 1 ALBUM)\n"
    "  MUSICIAN  1:N  ALBUM        [PRODUCES]    (each ALBUM has exactly 1 producer)"
)

add_heading(doc, "Assumptions and Non-capturable Constraints", level=2)
add_para(doc, "1. song_id is introduced as a surrogate key since song title may not be unique.")
add_para(doc, "2. album_identifier is treated as an alternate unique key.")
add_para(doc, "3. Constraint address -> phone is a functional dependency, generally enforced at schema/business-rule level.")

add_heading(doc, "Q2. Relational Algebra (Suppliers, Parts, Catalog)", level=1)
add_para(doc, "Given: Suppliers(sid, sname, address), Parts(pid, pname, color), Catalog(sid, pid, cost)")

add_heading(doc, "a) Names of suppliers who supply some red part", level=2)
add_code_block(doc, "π_sname ( Suppliers ⋈ Catalog ⋈ σ_color='red'(Parts) )")

add_heading(doc, "b) sids of suppliers who supply some red or green part", level=2)
add_code_block(doc, "π_sid ( Catalog ⋈ σ_color='red' OR color='green'(Parts) )")

add_heading(doc, "c) sids of suppliers who supply some red part or are at 221 Packer street", level=2)
add_code_block(
    doc,
    "π_sid ( Catalog ⋈ σ_color='red'(Parts) )\n"
    "UNION\n"
    "π_sid ( σ_address='221 Packer street'(Suppliers) )"
)

add_heading(doc, "d) sids of suppliers who supply some red part and some green part", level=2)
add_code_block(
    doc,
    "π_sid ( Catalog ⋈ σ_color='red'(Parts) )\n"
    "INTERSECT\n"
    "π_sid ( Catalog ⋈ σ_color='green'(Parts) )"
)

add_heading(doc, "e) sids of suppliers who supply every part", level=2)
add_code_block(doc, "π_sid,pid(Catalog) DIVIDE π_pid(Parts)")

add_heading(doc, "Q3. Relational Algebra and SQL", level=1)
add_para(doc, "Given: Flights, Aircraft, Certified, Employees")

add_heading(doc, "a) eids of employees who make the highest salary", level=2)
add_para(doc, "Relational Algebra:")
add_code_block(
    doc,
    "π_eid(Employees) - π_E1.eid( σ_E1.salary < E2.salary( ρ_E1(Employees) × ρ_E2(Employees) ) )"
)
add_para(doc, "SQL:")
add_code_block(
    doc,
    "SELECT e.eid\n"
    "FROM Employees e\n"
    "WHERE e.salary = (SELECT MAX(salary) FROM Employees);"
)

add_heading(doc, "b) eids of employees who make the second highest salary", level=2)
add_para(doc, "Relational Algebra: Not expressible in basic RA without aggregate/ranking support.")
add_para(doc, "SQL:")
add_code_block(
    doc,
    "SELECT e.eid\n"
    "FROM Employees e\n"
    "WHERE e.salary = (\n"
    "  SELECT MAX(salary)\n"
    "  FROM Employees\n"
    "  WHERE salary < (SELECT MAX(salary) FROM Employees)\n"
    ");"
)

add_heading(doc, "c) eids certified for the largest number of aircraft", level=2)
add_para(doc, "Relational Algebra: Not expressible in basic RA (requires grouped COUNT and MAX).")
add_para(doc, "SQL:")
add_code_block(
    doc,
    "WITH cnt AS (\n"
    "  SELECT eid, COUNT(*) AS c\n"
    "  FROM Certified\n"
    "  GROUP BY eid\n"
    ")\n"
    "SELECT eid\n"
    "FROM cnt\n"
    "WHERE c = (SELECT MAX(c) FROM cnt);"
)

add_heading(doc, "d) eids certified for exactly three aircraft", level=2)
add_para(doc, "Relational Algebra: Not expressible in basic RA (requires grouped COUNT).")
add_para(doc, "SQL:")
add_code_block(
    doc,
    "SELECT eid\n"
    "FROM Certified\n"
    "GROUP BY eid\n"
    "HAVING COUNT(*) = 3;"
)

add_heading(doc, "e) total amount paid to employees as salaries", level=2)
add_para(doc, "Relational Algebra: Not expressible in basic RA (requires SUM aggregate).")
add_para(doc, "SQL:")
add_code_block(
    doc,
    "SELECT SUM(salary) AS total_salary_paid\n"
    "FROM Employees;"
)

add_heading(doc, "Q4. SQL Queries (No duplicates in output)", level=1)
add_para(doc, "Given: Student, Class, Enrolled, Faculty")

add_heading(doc, "a) Faculty with combined enrollment less than five", level=2)
add_code_block(
    doc,
    "SELECT f.fname\n"
    "FROM Faculty f\n"
    "LEFT JOIN Class c ON c.fid = f.fid\n"
    "LEFT JOIN Enrolled e ON e.cname = c.name\n"
    "GROUP BY f.fid, f.fname\n"
    "HAVING COUNT(e.snum) < 5;"
)

add_heading(doc, "b) For each level, print level and average age", level=2)
add_code_block(
    doc,
    "SELECT s.level, AVG(s.age) AS avg_age\n"
    "FROM Student s\n"
    "GROUP BY s.level;"
)

add_heading(doc, "c) For all levels except JR, print level and average age", level=2)
add_code_block(
    doc,
    "SELECT s.level, AVG(s.age) AS avg_age\n"
    "FROM Student s\n"
    "WHERE s.level <> 'JR'\n"
    "GROUP BY s.level;"
)

add_heading(doc, "d) Faculty who taught classes only in room R128 with total classes taught", level=2)
add_code_block(
    doc,
    "SELECT f.fname, COUNT(*) AS total_classes\n"
    "FROM Faculty f\n"
    "JOIN Class c ON c.fid = f.fid\n"
    "GROUP BY f.fid, f.fname\n"
    "HAVING COUNT(*) > 0\n"
    "   AND SUM(CASE WHEN c.room <> 'R128' THEN 1 ELSE 0 END) = 0;"
)

add_heading(doc, "e) Names of students enrolled in maximum number of classes", level=2)
add_code_block(
    doc,
    "WITH scount AS (\n"
    "  SELECT e.snum, COUNT(*) AS cnt\n"
    "  FROM Enrolled e\n"
    "  GROUP BY e.snum\n"
    ")\n"
    "SELECT DISTINCT s.sname\n"
    "FROM Student s\n"
    "JOIN scount sc ON sc.snum = s.snum\n"
    "WHERE sc.cnt = (SELECT MAX(cnt) FROM scount);"
)

format_all_paragraphs(doc)

output_path = "DBMS-Assignment-01-Set-03-Solution-Submission-Final.docx"
doc.save(output_path)
print(output_path)
